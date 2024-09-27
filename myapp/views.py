from datetime import datetime  
from django.shortcuts import get_object_or_404, render, redirect
from .forms import DocumentForm
from .models import Document
import docx # type: ignore
import pdfplumber # type: ignore
import pytesseract # type: ignore
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from PIL import Image,ImageEnhance, ImageFilter
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_text_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_image(file):
    image = Image.open(file)
    
    # Convert the image to grayscale
    gray_image = image.convert('L')
    
    # Enhance the image contrast
    enhancer = ImageEnhance.Contrast(gray_image)
    enhanced_image = enhancer.enhance(2)
    
    # Apply a filter to remove noise
    filtered_image = enhanced_image.filter(ImageFilter.MedianFilter())
    
    # Perform OCR using Tesseract
    text = pytesseract.image_to_string(filtered_image, lang='eng', config='--psm 6')
    
    return text

@csrf_exempt
@login_required
def process_document(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.user = request.user
            document.created_at = datetime.now()  
            document = form.save()
            file_path = document.file.path
            file_extension = file_path.split('.')[-1].lower()

            if file_extension == 'pdf':
                text = extract_text_from_pdf(document.file)
            elif file_extension == 'docx':
                text = extract_text_from_docx(document.file)
            elif file_extension in ['jpg', 'jpeg', 'png']:
                text = extract_text_from_image(document.file)
            else:
                text = "Unsupported file format"

            document.converted_text = text
            document.save()
            

            return render(request, 'document_text.html', {'text': text, 'document_id': document.id})
    else:
        form = DocumentForm()
     # Show only documents belonging to the logged-in user
    documents = Document.objects.filter(user=request.user)
    context ={
            'form': form,
            'documents':documents     
        }

    return render(request, 'upload_document.html', context)

def view_document(request, document_id):
    # Ensure the document belongs to the logged-in user
    document = get_object_or_404(Document, id=document_id, user=request.user)
    return render(request, 'document_text.html', {'text': document.converted_text, 'document_id': document_id})



from django.http import HttpResponse
from reportlab.lib.pagesizes import letter # type: ignore
from reportlab.pdfgen import canvas # type: ignore
from docx import Document as DocxDocument # type: ignore
import io

# View to generate PDF

def download_as_pdf(request, document_id):
    document = Document.objects.get(id=document_id)
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    
    text = document.converted_text
    
    # Create a text object
    textobject = p.beginText()
    textobject.setTextOrigin(100, 750)
    textobject.setFont("Helvetica", 12)

    # Split the text into lines and add each line to the text object
    lines = text.split('\n')
    for line in lines:
        textobject.textLine(line)
    
    # Draw the text object on the canvas
    p.drawText(textobject)
    p.showPage()
    p.save()
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename=document_{document_id}.pdf'
    return response


# View to generate DOCX
def download_as_docx(request, document_id):
    document = Document.objects.get(id=document_id)
    doc = DocxDocument()
    doc.add_paragraph(document.converted_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=document_{document_id}.docx'
    return response

from django.contrib import messages

def delete_document(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    document.delete()
    messages.success(request, f'Document {document.file.name} was deleted successfully.')
    return redirect('process_document')

def demo(request,extension=None):
    return render(request,'404.html')

def home(request):
    return render(request,'login.html')

